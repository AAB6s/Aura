from pathlib import Path
from PIL import Image
import contextlib
import io
import json
import fitz
from docx import Document
from openpyxl import load_workbook
from .utils import normalize_text
from .config import OCR_ENABLED, PDF_RENDER_SCALE, LAYOUT_ENABLED, LAYOUT_MIN_AREA, LAYOUT_MAX_REGIONS_PER_PAGE

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from .layout_detector import get_layout_detector
except Exception:
    def get_layout_detector():
        return None

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp"}
TEXT_EXTS = {".txt", ".md", ".csv", ".tsv"}

def ocr_image(image: Image.Image) -> str:
    if not OCR_ENABLED or pytesseract is None:
        return ""
    try:
        return normalize_text(pytesseract.image_to_string(image))
    except Exception:
        return ""

def clamp_bbox(bbox, width, height):
    x1, y1, x2, y2 = [float(x) for x in bbox]
    x1 = max(0.0, min(float(width), x1))
    y1 = max(0.0, min(float(height), y1))
    x2 = max(0.0, min(float(width), x2))
    y2 = max(0.0, min(float(height), y2))
    if x2 < x1:
        x1, x2 = x2, x1
    if y2 < y1:
        y1, y2 = y2, y1
    return [round(x1, 2), round(y1, 2), round(x2, 2), round(y2, 2)]

def bbox_area(bbox):
    return max(0.0, bbox[2] - bbox[0]) * max(0.0, bbox[3] - bbox[1])

def bbox_iou(a, b):
    x1 = max(a[0], b[0])
    y1 = max(a[1], b[1])
    x2 = min(a[2], b[2])
    y2 = min(a[3], b[3])
    inter = max(0.0, x2 - x1) * max(0.0, y2 - y1)
    denom = bbox_area(a) + bbox_area(b) - inter
    return inter / denom if denom > 0 else 0.0

def crop_page_image(image, bbox, scale):
    px = [int(round(bbox[0] * scale)), int(round(bbox[1] * scale)), int(round(bbox[2] * scale)), int(round(bbox[3] * scale))]
    px[0] = max(0, min(image.width, px[0]))
    px[1] = max(0, min(image.height, px[1]))
    px[2] = max(0, min(image.width, px[2]))
    px[3] = max(0, min(image.height, px[3]))
    if px[2] <= px[0] or px[3] <= px[1]:
        return None
    return image.crop(tuple(px)).convert("RGB")

def block_text(block):
    parts = []
    for line in block.get("lines", []):
        line_parts = []
        for span in line.get("spans", []):
            text = span.get("text", "")
            if text:
                line_parts.append(text)
        if line_parts:
            parts.append(" ".join(line_parts))
    return normalize_text("\n".join(parts))

def normalize_layout_label(label):
    value = str(label or "unknown").strip().lower().replace(" ", "_").replace("-", "_")
    mapping = {
        "plain_text": "text",
        "paragraph": "text",
        "text_region": "text",
        "section_header": "title",
        "header": "title",
        "heading": "title",
        "picture": "image",
        "graphic": "image",
        "figure": "image",
        "photo": "image",
        "list_item": "list",
        "tabular": "table",
        "signature_stamp": "signature",
    }
    return mapping.get(value, value or "unknown")

def text_region_type(text, bbox, page_rect):
    low = text.lower()
    height = bbox[3] - bbox[1]
    width = bbox[2] - bbox[0]
    page_width = page_rect.width
    if "|" in text or "\t" in text:
        return "table"
    if any(x in low for x in ["signature", "signed", "stamp", "cachet"]):
        return "signature"
    if bbox[1] < page_rect.height * 0.18 and len(text) < 180 and height < page_rect.height * 0.12 and width > page_width * 0.25:
        return "title"
    if text.strip().startswith(("-", "*", "\u2022")):
        return "list"
    return "text"

def table_regions(page, page_number, image, scale, used_bboxes):
    rows = []
    try:
        with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
            tables = page.find_tables()
        items = getattr(tables, "tables", []) or []
    except Exception:
        items = []
    for idx, table in enumerate(items):
        bbox = clamp_bbox(table.bbox, page.rect.width, page.rect.height)
        if bbox_area(bbox) < LAYOUT_MIN_AREA:
            continue
        if any(bbox_iou(bbox, old) > 0.75 for old in used_bboxes):
            continue
        text_source = "none"
        try:
            extracted = table.extract()
            text = "\n".join(" | ".join(str(cell or "").strip() for cell in row if str(cell or "").strip()) for row in extracted)
            if normalize_text(text):
                text_source = "table_extractor"
        except Exception:
            text = ""
        crop = crop_page_image(image, bbox, scale)
        if not normalize_text(text) and crop is not None:
            text = ocr_image(crop)
            text_source = "ocr" if text else "none"
        rows.append({
            "text": normalize_text(text),
            "page": page_number,
            "image": crop,
            "kind": "layout_region",
            "region_type": "table",
            "bbox": bbox,
            "layout_confidence": 1.0,
            "element_id": f"p{page_number}_table_{idx + 1}",
            "text_source": text_source,
            "reading_order": len(rows) + 1,
            "page_width": round(page.rect.width, 2),
            "page_height": round(page.rect.height, 2),
        })
        used_bboxes.append(bbox)
    return rows

def native_layout_regions(page, page_number, image, scale, used_bboxes):
    rows = []
    try:
        data = page.get_text("dict")
        blocks = data.get("blocks", [])
    except Exception:
        blocks = []
    for block in blocks:
        bbox = clamp_bbox(block.get("bbox", [0, 0, 0, 0]), page.rect.width, page.rect.height)
        if bbox_area(bbox) < LAYOUT_MIN_AREA:
            continue
        if any(bbox_iou(bbox, old) > 0.70 for old in used_bboxes):
            continue
        btype = int(block.get("type", 0))
        crop = crop_page_image(image, bbox, scale)
        if btype == 1:
            text = ocr_image(crop) if crop is not None else ""
            region_type = "image"
            text_source = "ocr" if text else "none"
        else:
            text = block_text(block)
            region_type = text_region_type(text, bbox, page.rect)
            text_source = "pdf_native" if text else "none"
            if not text and crop is not None:
                text = ocr_image(crop)
                text_source = "ocr" if text else "none"
        rows.append({
            "text": normalize_text(text),
            "page": page_number,
            "image": crop,
            "kind": "layout_region",
            "region_type": region_type,
            "bbox": bbox,
            "layout_confidence": 1.0,
            "element_id": f"p{page_number}_r{len(rows) + 1}",
            "text_source": text_source,
            "reading_order": len(rows) + 1,
            "page_width": round(page.rect.width, 2),
            "page_height": round(page.rect.height, 2),
        })
        used_bboxes.append(bbox)
    return rows

def hf_layout_regions(page, page_number, image, scale, used_bboxes):
    detector = get_layout_detector() if LAYOUT_ENABLED else None
    if detector is None:
        return []
    rows = []
    try:
        detections = detector.detect(image)
    except Exception:
        detections = []
    for det in detections:
        px = det.get("bbox_px", [0, 0, 0, 0])
        bbox = clamp_bbox([px[0] / scale, px[1] / scale, px[2] / scale, px[3] / scale], page.rect.width, page.rect.height)
        if bbox_area(bbox) < LAYOUT_MIN_AREA:
            continue
        if any(bbox_iou(bbox, old) > 0.75 for old in used_bboxes):
            continue
        crop = crop_page_image(image, bbox, scale)
        region_type = normalize_layout_label(det.get("label"))
        text = normalize_text(page.get_textbox(fitz.Rect(*bbox)) or "")
        text_source = "pdf_native" if text else "none"
        if not text and crop is not None:
            text = ocr_image(crop)
            text_source = "ocr" if text else "none"
        rows.append({
            "text": normalize_text(text),
            "page": page_number,
            "image": crop,
            "kind": "layout_region",
            "region_type": region_type,
            "bbox": bbox,
            "layout_confidence": round(float(det.get("score", 0.0)), 5),
            "element_id": f"p{page_number}_hf_{len(rows) + 1}",
            "text_source": text_source,
            "reading_order": len(rows) + 1,
            "page_width": round(page.rect.width, 2),
            "page_height": round(page.rect.height, 2),
        })
        used_bboxes.append(bbox)
    return rows

def sort_regions(rows):
    rows = sorted(rows, key=lambda row: (row.get("bbox", [0, 0, 0, 0])[1], row.get("bbox", [0, 0, 0, 0])[0]))
    for idx, row in enumerate(rows, 1):
        row["reading_order"] = idx
        row["element_id"] = f"p{row.get('page')}_r{idx}"
    return rows[:LAYOUT_MAX_REGIONS_PER_PAGE]

def read_text_file(path: Path) -> list[dict]:
    for enc in ["utf-8", "utf-16", "latin-1"]:
        try:
            text = path.read_text(encoding=enc, errors="ignore")
            return [{"text": normalize_text(text), "page": None, "image": None, "kind": "text", "text_source": "file_text"}]
        except Exception:
            pass
    return []

def read_pdf(path: Path) -> list[dict]:
    rows = []
    doc = fitz.open(path)
    scale = PDF_RENDER_SCALE
    for i, page in enumerate(doc):
        page_number = i + 1
        text = normalize_text(page.get_text("text") or "")
        pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
        image = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
        if not text:
            text = ocr_image(image)
        overview = {
            "text": text,
            "page": page_number,
            "image": image,
            "kind": "pdf_page_overview",
            "region_type": "page",
            "bbox": [0.0, 0.0, round(page.rect.width, 2), round(page.rect.height, 2)],
            "layout_confidence": 1.0,
            "element_id": f"p{page_number}_page",
            "text_source": "pdf_native" if normalize_text(page.get_text("text") or "") else ("ocr" if text else "none"),
            "reading_order": 0,
            "page_width": round(page.rect.width, 2),
            "page_height": round(page.rect.height, 2),
        }
        rows.append(overview)
        used_bboxes = []
        regions = []
        regions.extend(table_regions(page, page_number, image, scale, used_bboxes))
        regions.extend(hf_layout_regions(page, page_number, image, scale, used_bboxes))
        regions.extend(native_layout_regions(page, page_number, image, scale, used_bboxes))
        regions = sort_regions(regions)
        if regions:
            rows.extend(regions)
    doc.close()
    return rows

def read_docx(path: Path) -> list[dict]:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
    return [{"text": normalize_text("\n".join(parts)), "page": None, "image": None, "kind": "docx", "text_source": "docx"}]

def read_xlsx(path: Path) -> list[dict]:
    wb = load_workbook(path, read_only=True, data_only=True)
    rows = []
    for ws in wb.worksheets:
        parts = []
        for row in ws.iter_rows(values_only=True):
            values = [str(v) for v in row if v not in [None, ""]]
            if values:
                parts.append(" | ".join(values))
        if parts:
            rows.append({"text": normalize_text("\n".join(parts)), "page": ws.title, "image": None, "kind": "xlsx_sheet", "text_source": "xlsx"})
    return rows

def recursive_json_text(obj, path="") -> list[dict]:
    rows = []
    if isinstance(obj, dict):
        page = obj.get("page", obj.get("pg"))
        element_id = str(obj.get("id", obj.get("element_id", path)))
        for key, value in obj.items():
            name = str(key).lower()
            if isinstance(value, str) and len(value.strip()) > 30 and any(x in name for x in ["text", "ocr", "caption", "content", "summary", "clause", "paragraph", "transcript"]):
                rows.append({"text": normalize_text(value), "page": page, "image": None, "kind": "json_text", "element_id": element_id, "field": key, "text_source": f"json.{key}"})
            elif isinstance(value, (dict, list)):
                rows.extend(recursive_json_text(value, f"{path}/{key}"))
    elif isinstance(obj, list):
        for i, value in enumerate(obj):
            rows.extend(recursive_json_text(value, f"{path}/{i}"))
    return rows

def read_json(path: Path) -> list[dict]:
    data = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
    return recursive_json_text(data)

def read_image(path: Path) -> list[dict]:
    image = Image.open(path).convert("RGB")
    text = ocr_image(image)
    return [{"text": text, "page": None, "image": image, "kind": "image", "region_type": "image", "bbox": [0.0, 0.0, float(image.width), float(image.height)], "element_id": "image_1", "text_source": "ocr" if text else "none", "reading_order": 0}]

def read_document(path: str | Path) -> list[dict]:
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        rows = read_pdf(path)
    elif ext == ".docx":
        rows = read_docx(path)
    elif ext in {".xlsx", ".xls"}:
        rows = read_xlsx(path)
    elif ext in TEXT_EXTS:
        rows = read_text_file(path)
    elif ext == ".json":
        rows = read_json(path)
    elif ext in IMAGE_EXTS:
        rows = read_image(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    for idx, row in enumerate(rows, 1):
        row.setdefault("source_file", path.name)
        row.setdefault("element_id", f"e{idx}")
        row.setdefault("region_type", row.get("kind", "document"))
        row.setdefault("bbox", None)
        row.setdefault("layout_confidence", None)
        row.setdefault("text_source", "unknown")
        row.setdefault("reading_order", idx)
        row["text"] = normalize_text(row.get("text", ""))
    return rows
